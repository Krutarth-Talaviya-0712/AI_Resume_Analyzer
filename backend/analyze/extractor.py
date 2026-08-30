"""
extractor.py
------------
Handles text extraction from uploaded resume files.
Supports PDF (via pdfplumber) and DOCX (via python-docx).
Returns plain text with whitespace normalized for downstream NLP.
"""

import io
import re

import pdfplumber
from docx import Document


def extract_text(file_storage) -> str:
    """
    Extract plain text from a werkzeug FileStorage object.

    Parameters
    ----------
    file_storage : werkzeug.datastructures.FileStorage
        The uploaded file object from Flask's request.files.

    Returns
    -------
    str
        Cleaned, normalized plain text from the resume.

    Raises
    ------
    ValueError
        If the file type is not supported (not PDF or DOCX).
    """
    filename = file_storage.filename.lower()

    if filename.endswith(".pdf"):
        return _extract_from_pdf(file_storage)
    elif filename.endswith(".docx"):
        return _extract_from_docx(file_storage)
    else:
        raise ValueError(
            "Unsupported file type. Please upload a PDF or DOCX resume."
        )


def _extract_from_pdf(file_storage) -> str:
    """Extract text from a PDF file using pdfplumber."""
    raw_bytes = file_storage.read()
    text_parts = []

    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return _clean_text("\n".join(text_parts))


def _extract_from_docx(file_storage) -> str:
    """Extract text from a DOCX file using python-docx."""
    raw_bytes = file_storage.read()
    document = Document(io.BytesIO(raw_bytes))

    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return _clean_text("\n".join(paragraphs))


def _clean_text(text: str) -> str:
    """
    Normalize extracted text:
    - Collapse multiple blank lines
    - Strip leading/trailing whitespace per line
    - Remove non-printable characters
    """
    # Remove non-printable characters except newlines and tabs
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E]", " ", text)
    # Normalize whitespace within lines
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    # Collapse more than two consecutive blank lines into one
    cleaned_lines = []
    blank_count = 0
    for line in lines:
        if not line:
            blank_count += 1
            if blank_count <= 1:
                cleaned_lines.append(line)
        else:
            blank_count = 0
            cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()
